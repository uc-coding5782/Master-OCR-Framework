"""DOCX exporter for OCR results."""

from __future__ import annotations

from pathlib import Path
from typing import TYPE_CHECKING

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from ocr_framework.exporters.base import Exporter
from ocr_framework.exceptions import ExportError
from ocr_framework.models.export_payload import ExportReport
from ocr_framework.models.page_result import DocumentResult

if TYPE_CHECKING:
    from ocr_framework.types import Metadata


class DOCXExporter(Exporter):
    """Export OCR results to DOCX format.

    The DOCXExporter produces Microsoft Word documents with recognized text,
    preserving page structure and formatting where possible.
    """

    @property
    def format(self) -> str:
        """Return the export format identifier."""
        return "docx"

    def export(
        self,
        result: DocumentResult,
        destination: Path,
        options: Metadata | None = None,
    ) -> ExportReport:
        """Export document result to a DOCX file.

        Args:
            result: OCR output to serialize.
            destination: Target output path.
            options: Optional exporter-specific settings.

        Returns:
            An ExportReport describing the export operation.

        Raises:
            ExportError: If export fails or python-docx is not available.
        """
        if not DOCX_AVAILABLE:
            raise ExportError(
                "python-docx is required for DOCX export. "
                "Install it with: pip install python-docx"
            )

        _ = options  # Reserved for future use

        try:
            # Ensure parent directory exists
            destination.parent.mkdir(parents=True, exist_ok=True)

            # Create Word document
            doc = DocxDocument()

            # Add content
            self._add_content(doc, result)

            # Save to file
            doc.save(str(destination))

            # Create export report
            report = ExportReport(
                format=self.format,
                destination=str(destination),
                success=True,
            )

            return report

        except IOError as exc:
            raise ExportError(f"Failed to write DOCX file: {exc}") from exc
        except Exception as exc:
            raise ExportError(f"DOCX export failed: {exc}") from exc

    def _add_content(self, doc: DocxDocument, result: DocumentResult) -> None:
        """Add OCR content to Word document.

        Args:
            doc: Word document object.
            result: Document result to add.
        """
        for page_result in result.pages:
            # Add page separator if multiple pages
            if len(result.pages) > 1:
                doc.add_heading(f"Page {page_result.page_index + 1}", level=2)

            # Add each line of text
            for line in page_result.lines:
                doc.add_paragraph(line.text)

            # Add page break between pages
            if len(result.pages) > 1 and page_result.page_index < len(result.pages) - 1:
                doc.add_page_break()
